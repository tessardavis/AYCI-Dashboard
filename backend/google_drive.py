"""
Google Drive + Docs integration (service-account) for private-tier doc summaries.

- Lists Google Docs / Shortcuts in a specific Drive folder.
- Fuzzy-matches filename against a student name.
- Exports the doc's plain-text content.
- Summarises via Emergent LLM (Claude Sonnet 4.5) to a ~200-word brief.

Results cached in Mongo `drive_doc_summaries` for 24 h so repeated Student
Lookups don't hit Google + LLM every time.
"""
from __future__ import annotations

import os
import re
import json
import logging
from datetime import datetime, timezone, timedelta
from difflib import SequenceMatcher
from typing import Optional

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

logger = logging.getLogger(__name__)

SCOPES = [
    "https://www.googleapis.com/auth/drive.readonly",
    "https://www.googleapis.com/auth/documents.readonly",
]

SUMMARY_TTL_HOURS = 24
MAX_DOC_CHARS = 12000  # Truncate input to the LLM

# Fuzzy-match thresholds tuned for typo tolerance (e.g. "Jevdovic" ↔ "Jedovic")
# without grabbing unrelated names. Anything ≥ FUZZY_MIN is surfaced with a
# verification flag for the coach.
FUZZY_MIN_WHOLE = 0.82
FUZZY_MIN_TOKEN = 0.78
# Below this, two name tokens are considered different people — a 5-char token
# pair under 0.55 is almost certainly unrelated.
FUZZY_HARD_FLOOR = 0.55


def _drive_service():
    """Build a Drive client from either a service-account JSON file path OR
    the raw JSON content pasted into the same env var. Lets non-technical
    deployers paste the JSON blob into Emergent Secrets without needing to
    add a separate filesystem path."""
    sa_value = (os.environ.get("GOOGLE_SERVICE_ACCOUNT_FILE") or "").strip()
    if not sa_value:
        raise RuntimeError("GOOGLE_SERVICE_ACCOUNT_FILE not configured")
    # If it's a file path that exists, load from disk
    if not sa_value.lstrip().startswith("{") and os.path.exists(sa_value):
        creds = service_account.Credentials.from_service_account_file(
            sa_value, scopes=SCOPES,
        )
    else:
        # Otherwise treat the value as raw JSON content
        try:
            info = json.loads(sa_value)
        except Exception as e:
            raise RuntimeError(
                f"GOOGLE_SERVICE_ACCOUNT_FILE is neither a valid path nor valid JSON: {e}"
            ) from e
        creds = service_account.Credentials.from_service_account_info(
            info, scopes=SCOPES,
        )
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def _folder_id() -> str:
    fid = os.environ.get("GOOGLE_DRIVE_PRIVATE_TIER_FOLDER_ID")
    if not fid:
        raise RuntimeError("GOOGLE_DRIVE_PRIVATE_TIER_FOLDER_ID not configured")
    return fid


def _normalise(name: str) -> str:
    """Lowercase + strip punctuation + collapse whitespace for fuzzy compare."""
    n = re.sub(r"[^a-z0-9\s]", "", name.lower())
    return re.sub(r"\s+", " ", n).strip()


def _ratio(a: str, b: str) -> float:
    return SequenceMatcher(None, a, b).ratio()


def _strip_extension(filename: str) -> str:
    """Strip a trailing .docx/.pdf/.doc/.gdoc/.txt/.md so it doesn't poison fuzzy ratio."""
    return re.sub(r"\.(docx?|pdf|gdoc|txt|md|rtf|odt)$", "", filename, flags=re.I)


def _token_match_score(target_tokens: list[str], file_tokens: list[str]) -> tuple[float, list[float]]:
    """
    For each target token, find its best fuzzy match against any file token.
    Returns (mean_score, per_token_scores). Tokens shorter than 3 chars
    (initials, middle name "M") are skipped — they shouldn't drive the score.
    """
    scores: list[float] = []
    for t in target_tokens:
        if len(t) < 3:
            continue
        best = 0.0
        for f in file_tokens:
            if len(f) < 3:
                continue
            r = _ratio(t, f)
            if r > best:
                best = r
        scores.append(best)
    if not scores:
        return 0.0, []
    return sum(scores) / len(scores), scores


def _find_best_match(student_name: str, files: list[dict]) -> Optional[dict]:
    """
    Match strategy (first win):
    1. Filename contains the full normalised student name (exact)
    2. All target tokens (>=3 chars) appear verbatim in filename (tokens)
    3. Fuzzy: highest SequenceMatcher ratio across whole-string OR per-token
       average — with a warning so the coach can verify
    4. Last-name verbatim fallback (lastname)

    Returns the matched file dict augmented with:
      - match_reason: "exact" | "tokens" | "fuzzy" | "lastname"
      - match_score: float (1.0 for strict matches, 0..1 for fuzzy)
      - needs_verification: True if match wasn't strict
      - other_candidates: top fuzzy near-misses for transparency (only for fuzzy)
    """
    target = _normalise(student_name)
    if not target:
        return None
    parts = [p for p in target.split() if p]

    # 1) exact substring
    for f in files:
        fname = _normalise(_strip_extension(f["name"]))
        if target in fname:
            return {**f, "match_reason": "exact", "match_score": 1.0, "needs_verification": False}

    # 2) all tokens present verbatim
    for f in files:
        fname = _normalise(_strip_extension(f["name"]))
        if all(p in fname for p in parts if len(p) >= 3):
            return {**f, "match_reason": "tokens", "match_score": 1.0, "needs_verification": False}

    # 3) fuzzy — score every file by whole-string AND per-token, take the max
    scored: list[tuple[float, str, dict, float, float]] = []
    for f in files:
        fname = _normalise(_strip_extension(f["name"]))
        whole = _ratio(target, fname)
        token_avg, _ = _token_match_score(parts, fname.split())
        score = max(whole, token_avg)
        scored.append((score, fname, f, whole, token_avg))
    scored.sort(key=lambda x: x[0], reverse=True)

    if scored:
        best_score, _, best_file, whole, tok = scored[0]
        # Only return as fuzzy if either the whole-string ratio OR the token
        # average crosses its threshold AND doesn't look like a totally different
        # name (token avg shouldn't be in the floor zone).
        passes = (whole >= FUZZY_MIN_WHOLE) or (tok >= FUZZY_MIN_TOKEN)
        if passes and tok >= FUZZY_HARD_FLOOR:
            others = [
                {"name": s[2]["name"], "score": round(s[0], 3)}
                for s in scored[1:4]
                if s[0] >= 0.65
            ]
            return {
                **best_file,
                "match_reason": "fuzzy",
                "match_score": round(best_score, 3),
                "needs_verification": True,
                "other_candidates": others,
            }

    # 4) fallback: last-name verbatim if it's an unusual surname
    if len(parts) >= 2 and len(parts[-1]) > 3:
        last = parts[-1]
        for f in files:
            fname = _normalise(_strip_extension(f["name"]))
            if last in fname.split():
                return {
                    **f,
                    "match_reason": "lastname",
                    "match_score": 1.0,
                    "needs_verification": True,
                }

    return None


async def _list_docs() -> list[dict]:
    """
    Returns list of {id, name, mimeType, modifiedTime, web_view_link,
    target_id (if shortcut), target_mime}.

    `supportsAllDrives` + `includeItemsFromAllDrives` enable Shared Drive folders.
    """
    drive = _drive_service()
    q = f"'{_folder_id()}' in parents and trashed=false"
    files: list[dict] = []
    page_token: Optional[str] = None
    while True:
        res = drive.files().list(
            q=q,
            pageSize=200,
            fields="nextPageToken, files(id,name,mimeType,modifiedTime,webViewLink,shortcutDetails)",
            pageToken=page_token,
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
            corpora="allDrives",
        ).execute()
        files.extend(res.get("files", []))
        page_token = res.get("nextPageToken")
        if not page_token:
            break
    out: list[dict] = []
    for f in files:
        item = {
            "id": f["id"],
            "name": f["name"],
            "mimeType": f["mimeType"],
            "modifiedTime": f.get("modifiedTime"),
            "web_view_link": f.get("webViewLink"),
        }
        if f["mimeType"] == "application/vnd.google-apps.shortcut":
            sc = f.get("shortcutDetails") or {}
            item["target_id"] = sc.get("targetId")
            item["target_mime"] = sc.get("targetMimeType")
        out.append(item)
    return out


# In-process cache of the Drive folder listing. The folder changes infrequently
# (a few new docs per week) so caching for 30 min eliminates 1-2s of latency
# on every cold drive-summary / find-link call.
_DOC_LIST_CACHE: dict = {"files": None, "fetched_at": None}
_DOC_LIST_TTL_SECONDS = 30 * 60


async def _list_docs_cached() -> list[dict]:
    """`_list_docs()` with a 30-min in-process cache. Use this for read paths
    where seeing a brand-new doc within 30 min isn't critical."""
    now = datetime.now(timezone.utc)
    fetched = _DOC_LIST_CACHE.get("fetched_at")
    if (
        _DOC_LIST_CACHE.get("files") is not None
        and fetched is not None
        and (now - fetched).total_seconds() < _DOC_LIST_TTL_SECONDS
    ):
        return _DOC_LIST_CACHE["files"]
    files = await _list_docs()
    _DOC_LIST_CACHE["files"] = files
    _DOC_LIST_CACHE["fetched_at"] = now
    return files


def _bust_doc_list_cache() -> None:
    _DOC_LIST_CACHE["files"] = None
    _DOC_LIST_CACHE["fetched_at"] = None


async def find_student_doc_link(db, student_name: str) -> dict:
    """
    Lightweight: returns just the Drive web view link for the matched private-tier
    doc (no AI summary, no body fetch). Cached for 24h per student name.
    Returns {found, web_view_link, name, match_reason}.
    """
    from datetime import datetime, timezone, timedelta
    if not student_name:
        return {"found": False, "web_view_link": None}
    key = f"drive_link:{_normalise(student_name)}"
    cutoff = datetime.now(timezone.utc) - timedelta(hours=24)
    cached = await db.cache.find_one({"_id": key}, {"_id": 0})
    if cached and cached.get("cached_at"):
        ca = cached["cached_at"]
        if ca.tzinfo is None:
            ca = ca.replace(tzinfo=timezone.utc)
        if ca > cutoff:
            return cached["payload"]
    try:
        files = await _list_docs_cached()
        match = _find_best_match(student_name, files)
        if not match:
            payload = {"found": False, "web_view_link": None, "name": None}
        else:
            payload = {
                "found": True,
                "web_view_link": match.get("web_view_link"),
                "name": match.get("name"),
                "match_reason": match.get("match_reason"),
                "match_score": match.get("match_score"),
                "needs_verification": match.get("needs_verification", False),
                "other_candidates": match.get("other_candidates", []),
            }
    except Exception as e:
        payload = {"found": False, "web_view_link": None, "error": str(e)}
    await db.cache.update_one(
        {"_id": key},
        {"$set": {"payload": payload, "cached_at": datetime.now(timezone.utc)}},
        upsert=True,
    )
    return payload


async def _fetch_doc_text(file_id: str, mime: str) -> str:
    """
    Read the file as plain text. Handles:
    - Google Docs (export to text/plain)
    - .docx (export via Drive's auto-conversion)
    - .html (download raw, strip tags)
    - .txt / .md (download raw)
    Anything else raises so the caller can show a friendly message.
    """
    drive = _drive_service()
    try:
        if mime == "application/vnd.google-apps.document":
            result = drive.files().export(
                fileId=file_id, mimeType="text/plain",
            ).execute()
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
        ):
            # Drive can convert Word docs on the fly
            result = drive.files().export_media(
                fileId=file_id, mimeType="text/plain",
            ).execute() if False else _docx_to_text(drive, file_id)
        elif mime in ("text/html", "application/xhtml+xml"):
            raw = drive.files().get_media(
                fileId=file_id, supportsAllDrives=True,
            ).execute()
            text = raw.decode("utf-8", errors="replace") if isinstance(raw, bytes) else str(raw)
            return _strip_html(text)
        elif mime in ("text/plain", "text/markdown"):
            raw = drive.files().get_media(
                fileId=file_id, supportsAllDrives=True,
            ).execute()
            return raw.decode("utf-8", errors="replace") if isinstance(raw, bytes) else str(raw)
        else:
            raise RuntimeError(f"Unsupported mime type for summarisation: {mime}")

        if isinstance(result, bytes):
            return result.decode("utf-8", errors="replace")
        return str(result)
    except HttpError as e:
        raise RuntimeError(f"Drive read failed: {e}") from e


def _docx_to_text(drive, file_id: str) -> str:
    """Download a .docx and return its plain text using python-docx."""
    raw = drive.files().get_media(fileId=file_id, supportsAllDrives=True).execute()
    if not isinstance(raw, bytes):
        raw = bytes(raw)
    try:
        import io
        from docx import Document  # python-docx
        doc = Document(io.BytesIO(raw))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except ImportError:
        raise RuntimeError("python-docx not installed — can't read .docx files")


def _strip_html(html: str) -> str:
    """Best-effort plaintext from an HTML blob."""
    import re
    # Remove script/style first
    html = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", html, flags=re.S | re.I)
    # Replace <br>, </p>, </div>, </li> with newlines for readability
    html = re.sub(r"<\s*(br|/p|/div|/li|/h[1-6])\s*/?>", "\n", html, flags=re.I)
    # Strip remaining tags
    html = re.sub(r"<[^>]+>", " ", html)
    # Decode common HTML entities
    html = (html.replace("&nbsp;", " ").replace("&amp;", "&")
                .replace("&lt;", "<").replace("&gt;", ">")
                .replace("&quot;", '"').replace("&#39;", "'"))
    # Collapse whitespace but keep paragraph breaks
    lines = [re.sub(r"\s+", " ", ln).strip() for ln in html.splitlines()]
    return "\n".join([ln for ln in lines if ln])


async def _summarise(doc_name: str, doc_text: str) -> str:
    """Use Claude to summarise the private-tier doc."""
    from llm_client import get_client, complete
    if get_client() is None:
        raise RuntimeError("ANTHROPIC_API_KEY not configured")

    text = doc_text[:MAX_DOC_CHARS]
    prompt = (
        "You are summarising an internal private-tier student plan document so the "
        "AYCI team can get up to speed before a call. Produce a concise "
        "briefing in plain English with short labelled sections. Avoid verbosity.\n\n"
        f"Document title: {doc_name}\n\n"
        "Document content:\n"
        f"---\n{text}\n---\n\n"
        "Return a briefing with these sections (only include a section if the doc covers it):\n"
        "- **Goals**: 1-2 lines on their interview/career goal\n"
        "- **Interview date(s)**: specific date(s) if mentioned\n"
        "- **Specialty & context**: specialty, hospital, any special context\n"
        "- **Progress so far**: calls done, mocks done, videos submitted, any notable wins/blockers\n"
        "- **Next steps**: what the team should focus on next\n"
        "- **Flags**: anything urgent or risks (e.g. confidence issues, cancelled slots, exam results)\n\n"
        "Keep the whole brief under 220 words."
    )

    response = await complete(
        system="Concise, factual briefing writer.",
        user=prompt,
        max_tokens=600,
    )
    return response.strip()


async def summarise_student_doc(db, student_name: str, student_email: str) -> dict:
    """
    Top-level API. Returns:
    {
      "found": bool,
      "file": {id, name, web_view_link, modifiedTime} | None,
      "summary": str | None,
      "cached": bool,
      "error": str | None,
      "candidates_scanned": int,
    }
    """
    if not student_name:
        return {"found": False, "file": None, "summary": None, "cached": False, "error": "No student name", "candidates_scanned": 0}

    cache_key = (student_email or "").strip().lower() or _normalise(student_name)
    try:
        # Check cache first
        cutoff = datetime.now(timezone.utc) - timedelta(hours=SUMMARY_TTL_HOURS)
        cached = await db.drive_doc_summaries.find_one({"_id": cache_key}, {"_id": 0})
        if cached:
            cached_at = cached.get("cached_at")
            if isinstance(cached_at, datetime):
                if cached_at.tzinfo is None:
                    cached_at = cached_at.replace(tzinfo=timezone.utc)
                if cached_at > cutoff:
                    return {**cached, "cached": True}

        files = await _list_docs_cached()
        match = _find_best_match(student_name, files)
        if not match:
            result = {
                "found": False,
                "file": None,
                "summary": None,
                "cached": False,
                "error": None,
                "candidates_scanned": len(files),
            }
            return result

        doc_id = match.get("target_id") or match["id"]
        target_mime = match.get("target_mime") or match["mimeType"]
        web_view_link = match.get("web_view_link")
        if not web_view_link and match.get("target_id"):
            web_view_link = f"https://drive.google.com/file/d/{match['target_id']}/view"

        SUPPORTED_MIMES = {
            "application/vnd.google-apps.document",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/html",
            "application/xhtml+xml",
            "text/plain",
            "text/markdown",
        }

        if target_mime not in SUPPORTED_MIMES:
            payload = {
                "found": True,
                "file": {
                    "id": doc_id,
                    "name": match["name"],
                    "web_view_link": web_view_link,
                    "modifiedTime": match.get("modifiedTime"),
                    "mime": target_mime,
                },
                "summary": None,
                "error": f"File type ({target_mime.split('.')[-1]}) not summarisable yet — open the link to view.",
                "candidates_scanned": len(files),
                "cached": False,
            }
        else:
            try:
                text = await _fetch_doc_text(doc_id, target_mime)
            except RuntimeError as exc:
                msg = str(exc)
                hint = (
                    "Couldn't read the doc — looks like the target file isn't "
                    "shared with the service account. Share it (or its parent "
                    "folder) with "
                    "ayci-drive-reader@ayci-dashboard.iam.gserviceaccount.com "
                    "(Viewer access) to enable the summary."
                ) if "notFound" in msg or "not found" in msg.lower() or "403" in msg else msg
                return {
                    "found": True,
                    "file": {
                        "id": doc_id,
                        "name": match["name"],
                        "web_view_link": web_view_link,
                        "modifiedTime": match.get("modifiedTime"),
                        "mime": target_mime,
                    },
                    "summary": None,
                    "error": hint,
                    "candidates_scanned": len(files),
                    "cached": False,
                    "match": {
                        "reason": match.get("match_reason"),
                        "score": match.get("match_score"),
                        "needs_verification": match.get("needs_verification", False),
                        "other_candidates": match.get("other_candidates", []),
                        "searched_name": student_name,
                    },
                }
            if not text.strip():
                summary = "Document is empty."
            else:
                summary = await _summarise(match["name"], text)
            payload = {
                "found": True,
                "file": {
                    "id": doc_id,
                    "name": match["name"],
                    "web_view_link": web_view_link,
                    "modifiedTime": match.get("modifiedTime"),
                    "mime": target_mime,
                },
                "summary": summary,
                "error": None,
                "candidates_scanned": len(files),
                "cached": False,
            }

        # Store in cache
        # Stitch on match metadata so the UI can flag fuzzy matches.
        payload["match"] = {
            "reason": match.get("match_reason"),
            "score": match.get("match_score"),
            "needs_verification": match.get("needs_verification", False),
            "other_candidates": match.get("other_candidates", []),
            "searched_name": student_name,
        }
        await db.drive_doc_summaries.update_one(
            {"_id": cache_key},
            {"$set": {**payload, "cached_at": datetime.now(timezone.utc)}},
            upsert=True,
        )
        return payload
    except Exception as e:
        logger.exception(f"summarise_student_doc failed: {e}")
        # Don't cache errors — let the next attempt retry
        return {
            "found": False,
            "file": None,
            "summary": None,
            "cached": False,
            "error": str(e),
            "candidates_scanned": 0,
        }
